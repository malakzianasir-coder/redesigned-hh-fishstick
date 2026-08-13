"""Build a non-technical Word brief on website messaging changes for hospital management."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x24, 0x41)
RED = RGBColor(0xC4, 0x00, 0x12)
MUTED = RGBColor(0x5B, 0x61, 0x73)
RULE = "D8DCE6"
HEADER_FILL = "1B2441"
ALT_FILL = "F4F6FA"

OUT = Path(r"C:\Work\hijaz-2026\Hijaz-Hospital-Website-Messaging-Changes.docx")


def set_run_font(run, *, name="Calibri", size=11, bold=False, italic=False, color=NAVY):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color=RULE, sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, size=10, color=NAVY, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = 1


def add_para(doc, text, *, size=11, bold=False, italic=False, color=NAVY, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, *, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=NAVY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, color=NAVY)
    return p


def add_example_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.autofit = True
    headers = ("Page / area", "Before (what visitors saw)", "After (what visitors see now)")
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
    for r_i, row in enumerate(rows):
        fill = ALT_FILL if r_i % 2 else None
        for c_i, val in enumerate(row):
            set_cell_text(table.rows[r_i + 1].cells[c_i], val, size=9, fill=fill)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Cover / header
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("Hijaz Hospital Trust")
    set_run_font(r, size=12, bold=True, color=RED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Website Messaging Updates")
    set_run_font(r, size=22, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run("A plain-language summary of content and wording changes")
    set_run_font(r, size=12, italic=True, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    r = meta.add_run(
        "Prepared for hospital leadership  ·  13 August 2026\n"
        "Prepared by NorthWest Innovations Labs / TwoCX LLC"
    )
    set_run_font(r, size=10, color=MUTED)

    add_para(
        doc,
        "This note explains, in everyday language, what changed on the public website’s "
        "wording — titles, short introductions, and section labels — and why. It is not a "
        "technical document. Page names (the main titles at the top of each page) were kept "
        "as they were.",
    )

    # Why
    add_heading(doc, "1. Why we made these changes")
    add_para(
        doc,
        "As the website grew, the same sentence often appeared two or three times on one page — "
        "once as a small red label, again as the section title, and again as the short paragraph "
        "under the main heading. Visitors and donors should not have to read the same line repeatedly. "
        "We also wanted every page to sound like one hospital: clear, dignified, and compassionate.",
    )
    add_bullet(doc, "Remove repeated wording on the same page.")
    add_bullet(doc, "Keep a consistent voice across About, Patient Care, Donate, and Welfare pages.")
    add_bullet(doc, "Keep small labels useful for orientation (“where am I?”) without copying the page title.")
    add_bullet(doc, "Keep main page titles unchanged so bookmarks, printouts, and staff habits stay familiar.")

    # What we did not change
    add_heading(doc, "2. What we deliberately did not change")
    add_bullet(doc, "Main page titles (the large heading at the top of each page) — left as agreed.")
    add_bullet(doc, "Doctors directory wording — not part of this pass.")
    add_bullet(doc, "Lab tests catalogue wording — not part of this pass.")
    add_bullet(doc, "Long historical narratives (for example the founding journey) — only fixed where the same sentence was duplicated on the page.")
    add_bullet(
        doc,
        "About Us filter labels (Our Purpose, Leadership & Governance, Messages, Our Impact) — kept as the official page names visitors expect.",
    )

    # Guiding idea
    add_heading(doc, "3. The guiding message")
    add_para(
        doc,
        "Care within reach. Hijaz Hospital Trust is a non-profit, trust-based hospital where "
        "healthcare is treated as a right, not a privilege — delivered with compassion, dignity, "
        "and clarity. We use that idea across the site without pasting the identical sentence on every page.",
    )

    # What changed by area
    add_heading(doc, "4. What changed, by area of the website")

    add_para(doc, "About the hospital (Purpose, Leadership, Messages, Impact, Supporters)", bold=True, space_after=4)
    add_para(
        doc,
        "Short introductions under the main titles were rewritten so they no longer copy the "
        "search-engine description word-for-word. Leadership and chairman/president message pages "
        "now open with a clearer one-line summary of what the visitor is about to read.",
    )

    add_para(doc, "Donate and ways to give", bold=True, space_after=4)
    add_para(
        doc,
        "On the main Donate page, the Qur’anic verse remains the single supporting line under "
        "“Ways to Give.” The long paragraph that repeated the same idea underneath was removed "
        "so the verse can stand with dignity. Individual giving pages (Zakat, Sadaqah, meals, "
        "sponsorships, and so on) keep their campaign lines on the page, while behind-the-scenes "
        "search descriptions were clarified so they are not identical copies.",
    )

    add_para(doc, "Patient Care (hospital services such as OPD, IPD, ICU, Emergency, Laboratory)", bold=True, space_after=4)
    add_para(
        doc,
        "Many service pages previously showed two almost identical lines under the title — a short "
        "slogan and then the same idea again as a description. Those are now merged into one clear "
        "supporting sentence. All Patient Care service pages also use the same refined presentation "
        "style for that sentence (the italic “quote-style” introduction already used on department pages), "
        "so the family of pages feels consistent.",
    )

    add_para(doc, "Clinical departments", bold=True, space_after=4)
    add_para(
        doc,
        "Department page names and clinical facts were preserved. Search descriptions were separated "
        "from the on-page introduction so Google and the page itself are not forced to show the exact "
        "same sentence.",
    )

    add_para(doc, "Patient welfare programmes", bold=True, space_after=4)
    add_para(
        doc,
        "Programme names stayed the same. Short meta descriptions were adjusted so they do not "
        "duplicate the visible introduction on each welfare page.",
    )

    # Examples
    add_heading(doc, "5. Plain examples (before → after)")
    add_para(
        doc,
        "These examples show the kind of change a visitor would notice. They are representative, not an exhaustive list.",
        italic=True,
        color=MUTED,
        space_after=8,
    )

    add_example_table(
        doc,
        [
            (
                "Donate (Ways to Give)",
                "Qur’an verse plus a long paragraph repeating the call to give",
                "Qur’an verse alone as the supporting line under the title",
            ),
            (
                "Inpatient Department (IPD)",
                "Slogan and nearly identical description stacked under the title",
                "One clear supporting sentence; same page title",
            ),
            (
                "Our Purpose",
                "Introduction copied the long meta description",
                "Shorter on-page line about vision, journey, and values — care within reach",
            ),
            (
                "Chairman’s Message",
                "Generic “shares his vision…” line used twice",
                "One specific line: gratitude, resolve, and service",
            ),
            (
                "About Us filters",
                "(Briefly renamed during drafting)",
                "Restored to official names: Our Purpose, Leadership & Governance, Messages, Our Impact",
            ),
            (
                "Section labels (e.g. OPD Timings)",
                "Small red label and section title saying the same words",
                "Small label shows the role (e.g. Timings); title keeps the full name",
            ),
        ],
    )

    # Benefits
    add_heading(doc, "6. Benefits for the hospital")
    add_bullet(doc, "Clearer first impression — visitors are not asked to read the same sentence twice.")
    add_bullet(doc, "Stronger brand consistency — About, Care, Donate, and Welfare sound like one institution.")
    add_bullet(doc, "Better orientation — small labels still tell people which part of the site they are in.")
    add_bullet(doc, "Familiar page names — staff and partners can still refer to pages by the titles they already know.")
    add_bullet(doc, "Easier future updates — a short content standard is in place so new pages follow the same rules.")

    # Review
    add_heading(doc, "7. Suggested review by leadership")
    add_para(
        doc,
        "We welcome confirmation on tone and priorities. In particular, please glance at:",
    )
    add_bullet(doc, "Donate — Qur’an verse standing alone under “Ways to Give.”")
    add_bullet(doc, "Any Patient Care service page (for example OPD or IPD) — single introduction line.")
    add_bullet(doc, "Our Purpose, Leadership, and leadership message pages — refreshed opening lines.")
    add_bullet(doc, "About Us — filter names matching the official section titles.")
    add_para(
        doc,
        "If any line should return to a preferred house wording, send the preferred text and we will align it.",
        space_after=12,
    )

    # Closing
    add_heading(doc, "8. Prepared by")
    add_para(
        doc,
        "NorthWest Innovations Labs / TwoCX LLC\n"
        "Website content and messaging pass for Hijaz Hospital Trust\n"
        "13 August 2026",
        space_after=4,
    )
    add_para(
        doc,
        "For questions about this summary or preferred alternate wording, contact your project lead at NorthWest Innovations Labs / TwoCX LLC.",
        italic=True,
        color=MUTED,
        size=10,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
