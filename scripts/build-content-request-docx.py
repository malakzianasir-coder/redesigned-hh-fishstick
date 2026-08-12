"""Build the Hijaz Hospital content-request Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x24, 0x41)
RED = RGBColor(0xC4, 0x00, 0x12)
MUTED = RGBColor(0x5B, 0x61, 0x73)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RULE = "D8DCE6"
HEADER_FILL = "1B2441"
ALT_FILL = "F4F6FA"
MUST_FILL = "FDECEC"
REPLACE_FILL = "FFF6E8"
CONFIRM_FILL = "EEF4FF"

OUT = Path(r"C:\Work\hijaz-2026\Hijaz-Hospital-Website-Content-and-Photo-Request.docx")


def set_run_font(run, *, name="Calibri", size=11, bold=False, italic=False, color=NAVY):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color=RULE, sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, size=10, color=NAVY, align="left", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = 1  # center


def set_table_widths(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def add_num_pages(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " NUMPAGES "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def para(doc, text, *, size=11, bold=False, italic=False, color=NAVY, space_after=8, space_before=0, align="left"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=NAVY)
    else:
        set_run_font(run, size=13, bold=True, color=NAVY)
    # underline bar via bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12" if level == 1 else "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C40012" if level == 1 else "D8DCE6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_table(doc, headers, rows, widths, header_fill=HEADER_FILL):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, color=WHITE, fill=header_fill)
    for r_i, row in enumerate(rows):
        fill = ALT_FILL if r_i % 2 else "FFFFFF"
        for c_i, val in enumerate(row):
            align = "center" if c_i == 0 and headers[0] in ("#",) else "left"
            set_cell_text(
                table.rows[r_i + 1].cells[c_i],
                str(val),
                size=10,
                color=NAVY,
                align=align,
                fill=fill,
                bold=c_i == 0,
            )
    set_table_widths(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead:
        r1 = p.add_run("•  " + bold_lead)
        set_run_font(r1, size=11, bold=True, color=NAVY)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=NAVY)
    else:
        r = p.add_run("•  " + text)
        set_run_font(r, size=11, color=NAVY)
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.7)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("NORTHWEST INNOVATIONS")
    set_run_font(r, size=11, bold=True, color=NAVY)
    r2 = hp.add_run("  /  TwoCX LLC")
    set_run_font(r2, size=11, bold=False, color=RED)
    hp2 = header.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r3 = hp2.add_run("Prepared for Hijaz Hospital  ·  Website content & photo request")
    set_run_font(r3, size=9, color=MUTED)
    pPr = hp2._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "C40012")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), "D8DCE6")
    pBdr.append(top)
    pPr.append(pBdr)
    r = fp.add_run("Confidential  ·  NorthWest Innovations / TwoCX LLC  ·  12 August 2026")
    set_run_font(r, size=9, color=MUTED)
    tab = fp.add_run("\t")
    set_run_font(tab, size=9, color=MUTED)
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(17.4), alignment=WD_TAB_ALIGNMENT.RIGHT)
    rpage = fp.add_run("Page ")
    set_run_font(rpage, size=9, color=MUTED)
    add_page_number(fp)
    rof = fp.add_run(" of ")
    set_run_font(rof, size=9, color=MUTED)
    add_num_pages(fp)

    # Cover / intro
    para(doc, "HIJAZ HOSPITAL", size=11, bold=True, color=RED, space_after=2)
    para(doc, "Website Content and Photo Request", size=22, bold=True, color=NAVY, space_after=6)
    para(
        doc,
        "A checklist of pictures, logos, and wording still needed so the public website can go live without placeholders, stock photos, or guesswork.",
        size=12,
        color=MUTED,
        space_after=12,
    )

    meta = doc.add_table(rows=2, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["Prepared by", "Prepared for", "Date", "Status"]
    values = [
        "NorthWest Innovations / TwoCX LLC",
        "Hijaz Hospital Trust",
        "12 August 2026",
        "Awaiting assets",
    ]
    for i, (lab, val) in enumerate(zip(labels, values)):
        set_cell_text(meta.rows[0].cells[i], lab, bold=True, size=8, color=WHITE, fill=HEADER_FILL, align="center")
        set_cell_text(meta.rows[1].cells[i], val, size=9, color=NAVY, fill=ALT_FILL, align="center")
    set_table_widths(meta, [4.35, 4.35, 4.35, 4.35])

    para(doc, "", space_after=6)
    para(
        doc,
        "Please return files by email or shared folder. JPEG or WebP is fine for photographs. Portraits work best as a clear head-and-shoulders shot on a plain background. Department and service photos should show the actual room, equipment, or activity named below — not a similar area of the hospital.",
        space_after=10,
    )

    heading(doc, "How to use this list", 1)
    para(doc, "Items are grouped by type. Each item says where it appears, what we need, and why.", space_after=8)
    add_table(
        doc,
        ["Priority", "Meaning"],
        [
            ["Must have", "Currently blank, a grey placeholder, or a QR/logo box with no real file"],
            ["Please replace", "Something is on the page, but it is the wrong subject, a stand-in, or a stock photo"],
            ["Please confirm / write", "Text, names, dates, or a decision — not a photo"],
        ],
        [4.5, 12.9],
    )

    # Section 1
    heading(doc, "1. Must have — missing pictures and files", 1)

    heading(doc, "1.1 Doctor portraits (grey “photo coming soon”)", 2)
    para(
        doc,
        "These consultants appear on Find a Doctor (hijazhospital.com/doctors) with no photograph.",
        space_after=8,
    )
    add_table(
        doc,
        ["#", "Name", "Role / department"],
        [
            ["1", "Dr. Hani Gul", "Anesthetist"],
            ["2", "Dr. Attique Abou Bakr", "Gastroenterologist (Head of Department)"],
            ["3", "Dr. Syeda Aneeqa Bano", "Gynecologist"],
            ["4", "Dr. Tasmeen Afridi", "Gynecologist"],
            ["5", "Dr. Fatima Zubair", "Gynecologist Consultant"],
            ["6", "Dr. Hafiza Sadia Jafar", "Jr. Consultant Anesthesia"],
            ["7", "Dr. Iffat Niazi", "Jr. Consultant Obs & Gynae"],
            ["8", "Dr. Saima Afzal", "Jr. Consultant Obs & Gynae"],
            ["9", "Dr. Uzma Hanif", "Jr. Consultant Obs & Gynae"],
            ["10", "Dr. Muhammad Hasan", "Junior Consultant Ortho"],
            ["11", "Dr. Ayesha Azhar", "Nephrologist (Head of Department)"],
            ["12", "Dr. Muhammad Amir Ishaq", "Pediatric Surgeon (Head of Department)"],
            ["13", "Dr. Amna Faisal", "Radiologist"],
            ["14", "Prof. Farooq Hameed", "Urologist"],
            ["15", "Dr. Nabiha Rizvi", "Nephrologist (visiting)"],
            ["16", "Dr. Muhammad Ali", "Senior Registrar Medicine, FMH (visiting)"],
            ["17", "Dr. Fatima Kamran", "Senior Registrar General Surgery, FMH (visiting)"],
            ["18", "Dr. Anum", "Senior Registrar Ophthalmology, FMH (visiting) — please also confirm full name"],
            ["19", "Dr. Unsia Rao", "Senior Registrar Ophthalmology, FMH (visiting)"],
        ],
        [1.4, 5.8, 10.2],
    )

    para(
        doc,
        "Please also send portraits for the homepage “Meet Our Expert Medical Team” block:",
        space_after=8,
    )
    add_table(
        doc,
        ["#", "Name", "Shown as", "Note"],
        [
            ["20", "Dr. Arfan Ishaque", "Medical Superintendent", "Grey placeholder on home"],
            ["21", "Dr. Muhammad Iqbal Hussain", "Visiting Consultant", "Grey placeholder — is this person on staff? See §3.2"],
            ["22", "Dr. Sana Malik", "Consultant Gynaecologist", "Grey placeholder — not in the doctors directory"],
            ["23", "Dr. Imran Shah", "Consultant Surgeon", "Grey placeholder — not in the doctors directory"],
            ["24", "Dr. Ayesha Khan", "Consultant Paediatrician", "Grey placeholder — not in the doctors directory"],
        ],
        [1.4, 5.2, 4.6, 6.2],
    )
    para(
        doc,
        "Dr. Nadeem Iqbal already has a photo. Please confirm whether he should still appear on the homepage team, given the news item about his farewell after 15 years.",
        space_after=10,
    )

    heading(doc, "1.2 QR codes for donations", 2)
    para(doc, "These pages currently show an empty QR box. Please send the official QR image files (PNG is fine).", space_after=8)
    add_table(
        doc,
        ["#", "What", "Where it appears"],
        [
            ["25", "JazzCash donation QR", "Home “Ways to Give” and How to Donate → Mobile Wallet"],
            ["26", "Meezan / donation QR", "How to Donate → Meezan App & QR"],
        ],
        [1.4, 5.5, 10.5],
    )

    heading(doc, "1.3 Logos still missing", 2)
    add_table(
        doc,
        ["#", "What", "Where it appears"],
        [
            ["27", "Sehat Sahulat Program logo", "IPD page (government coverage box) and Patient Welfare → Sehat Sahulat"],
            ["28", "BankIslami logo", "Bank details for the Medical Tower construction account (home, Donate, Bank Transfer)"],
            ["29", "Certificate of Incorporation (logo or scan)", "Our Purpose → compliance section — title present, no logo"],
            ["30", "National Distributors logo", "Our Supporters donor wall, #13"],
            ["31", "Eiffel Industries (Pvt.) Limited logo", "Our Supporters donor wall, #19"],
        ],
        [1.4, 6.2, 9.8],
    )
    para(
        doc,
        "Optional but useful: website addresses for National Distributors, Eiffel Industries, Northstar Textile, Astron Chemicals, and Masha Trading Co. (China) if you want them linked.",
        italic=True,
        color=MUTED,
        space_after=10,
    )

    heading(doc, "1.4 Room and facility photos that do not exist yet", 2)
    add_table(
        doc,
        ["#", "What", "Where it appears"],
        [
            ["32", "Cafeteria / dining area", "Services → Cafeteria — currently a generic hospital photo, not the cafeteria"],
            ["33", "Private Gynecology room", "Services → IPD → Accommodation (4 rooms)"],
            ["34", "VIP room", "Same page (2 rooms)"],
            ["35", "Semi-VIP room", "Same page (1 room)"],
            ["36", "VVIP room", "Same page (1 room)"],
        ],
        [1.4, 5.8, 10.2],
    )
    para(doc, "The accommodation section currently shows empty photo boxes next to the room counts.", space_after=10)

    heading(doc, "1.5 Leadership biographies still without their own pages", 2)
    para(
        doc,
        "We have portraits and messages, but two full life-story pages from the original Profiles document were never built:",
        space_after=8,
    )
    add_table(
        doc,
        ["#", "Person", "What we need"],
        [
            [
                "37",
                "Mrs. Tasneem Firdous Waheed",
                "Birth and death years (source still says “19XX–20XX”), plus confirmation that the existing biography text is approved",
            ],
            [
                "38",
                "Sohail Iqbal (Vohra?)",
                "Confirm whether the surname “Vohra” should appear on the site; confirm the Profiles biography is approved for a dedicated page",
            ],
        ],
        [1.4, 5.5, 10.5],
    )

    # Section 2
    heading(doc, "2. Please replace — wrong, stand-in, or stock pictures", 1)
    para(
        doc,
        "These pages already have an image, but it is not a true photo of that subject. Visitors will notice.",
        space_after=8,
    )

    heading(doc, "2.1 Department pages using a photo from another area", 2)
    add_table(
        doc,
        ["#", "Page", "Current problem"],
        [
            ["39", "Orthopedics", "Dental-unit photo used as stand-in"],
            ["40", "Urology", "Dialysis-ward photo used as stand-in"],
            ["41", "Plastic & Reconstructive Surgery", "Operating-theatre photo reused from General Surgery"],
            ["42", "ENT (Ear, Nose & Throat)", "Generic corridor / department photo"],
            ["43", "Gastroenterology", "Generic department photo"],
            ["44", "Pulmonology", "Male-ward photo used as stand-in"],
            ["45", "Endocrinology", "Generic department photo"],
            ["46", "Dermatology", "Generic department photo"],
            ["47", "Dietetics & Nutrition", "Waiting-area photo used as stand-in"],
        ],
        [1.4, 6.5, 9.5],
    )

    heading(doc, "2.2 Service pages", 2)
    add_table(
        doc,
        ["#", "Page", "Current problem"],
        [
            ["48", "Ambulance", "ICU photo used as stand-in — we still need a real ambulance / ambulance-bay photograph"],
            ["49", "Anesthesia", "ICU photo used as stand-in"],
            ["50", "Cafeteria", "Same as #32 — not a cafeteria photo"],
        ],
        [1.4, 4.5, 11.5],
    )

    heading(doc, "2.3 Homepage machinery carousel", 2)
    para(
        doc,
        "The “Latest Machinery & Advanced Medical Technology” strip on the homepage currently reuses ward and service banners (dialysis unit, ICU, lab, etc.), not close-up photos of named machines.",
        space_after=6,
    )
    para(doc, "Please send labelled photos of the actual equipment you want featured, for example:", space_after=4)
    for item in [
        "Dialysis machine",
        "Ultrasound",
        "Digital X-ray",
        "ICU monitors",
        "Cardiac diagnostics equipment",
        "Any new machines (Lasotronix or others) you want named on the site",
    ]:
        bullet(doc, item)

    heading(doc, "2.4 News, events, and success stories — stock or mismatched photos", 2)
    para(
        doc,
        "Several news and event stories use a photo from a different occasion (for example a dengue walk using a paediatrics opening photo, Nurses Day using a dialysis photo). Some success stories and event cards use Pexels stock photography, which should not stay on a live hospital site.",
        space_after=8,
    )
    para(
        doc,
        "Please send one real photograph per story, taken at that event or of that patient’s care setting (with consent).",
        space_after=8,
    )
    para(doc, "News / events that need their own photos", bold=True, space_after=6)
    add_table(
        doc,
        ["Story", "Where it appears"],
        [
            ["Dengue Awareness Campaign / Walk 2026", "News, Events, homepage news"],
            [
                "Annual Farewell Dinner 2026",
                "News, Events (one farewell photo already exists — please confirm it is the right event)",
            ],
            ["International Nurses Day 2026", "News, Events"],
            ["Visit of Dr. Muhammad Iqbal Hussain", "News"],
            ["Farewell ceremony for Dr. Nadeem Iqbal", "News"],
            ["Other news items currently sharing ICU / gynae / visitors / dialysis photos", "News"],
        ],
        [7.5, 9.9],
    )
    para(doc, "Success stories (Success Stories page and homepage)", bold=True, space_after=6)
    para(
        doc,
        "The nine published stories are sample write-ups (generic titles such as “Patient’s son”, “Patient”). They also mix hospital photos with stock images. Please either:",
        space_after=4,
    )
    bullet(doc, "approve named, consented patient stories with real photos, or")
    bullet(doc, "tell us to take these sample stories off the public site until real ones are ready.")
    para(
        doc,
        "Homepage success-story cards still say “Patient story title” with generic body copy. They need real titles and short summaries once stories are approved.",
        space_before=6,
        space_after=6,
    )
    para(
        doc,
        "Department pages also show a dashed box noting that stories will appear once tagged. That is expected until real stories exist.",
        italic=True,
        color=MUTED,
        space_after=10,
    )

    # Section 3
    heading(doc, "3. Please write or confirm — text, names, and facts", 1)

    heading(doc, "3.1 Homepage wording that still looks unfinished", 2)
    add_table(
        doc,
        ["#", "What", "Where"],
        [
            ["51", "Button currently reads “View repots” — should this be “View reports”?", "Home, “How You Can Engage”"],
            ["52", "Button currently reads “Complaint now” — preferred wording? (e.g. “Submit a complaint”)", "Same block"],
            ["53", "Three success-story cards titled “Patient story title”", "Home, Success Stories"],
            [
                "54",
                "Event blurbs that still mention the source file (“from merged.txt → Our Impact.txt”)",
                "Home events: Youm-e-Pakistan, Nursing Day — please supply a short public sentence for each",
            ],
        ],
        [1.4, 8.5, 7.5],
    )

    heading(doc, "3.2 Names, spelling, and dates to confirm", 2)
    add_table(
        doc,
        ["#", "Question"],
        [
            ["55", "Founder’s name: Inam Ellahi Asar vs Inam Elahi Asar (and “Haji” / “Hahi” in older notes). Which spelling is official?"],
            ["56", "President: Sohail Iqbal vs Sohail Iqbal Vohra — which should appear on the site?"],
            ["57", "Mrs. Tasneem Firdous Waheed — birth year and death year"],
            ["58", "Chairperson tenure for Mrs. Tasneem: source once said “20110 to 2020”. We show 2010–2020. Please confirm."],
            ["59", "Dr. Attique Abou Bakr is listed as Gastroenterologist but currently grouped under ENT in the directory. Please confirm the correct department."],
            ["60", "Dr. Muhammad Afaq appears twice (consultant and Head of Orthopedics). Keep one card or both?"],
            ["61", "Vacant — Medical Specialist (Head of Department) is in the original doctors list but hidden on the site. Should a “Position vacant” card appear?"],
            ["62", "Homepage team: should Dr. Muhammad Iqbal Hussain (US visitor) and Dr. Nadeem Iqbal (farewell) be listed as current team? If not, which six consultants should appear?"],
            ["63", "Dr. Anum (FMH Ophthalmology) — full name?"],
        ],
        [1.4, 16.0],
    )

    heading(doc, "3.3 Content that was never finished in the original documents", 2)
    add_table(
        doc,
        ["#", "Topic", "What we need"],
        [
            ["64", "Hospital Achievements", "Original note: “Share success stories, need to review yet.” Please send an approved achievements list, or say this section should stay off the site."],
            ["65", "Future Roadmap", "MRI, CT, OCT, YAG Laser, Plastic Surgery, Cardiac Surgery, Psychiatry, Chemo/Radiotherapy — is this still the plan, and should it appear on Our Impact?"],
            ["66", "“Our Kind Souls”", "Seven names from an older presentation (including Salah ud Din A. Sahaf, Sh. Azeem Pasha, Majeed A. Khan, Sh. Atta ur Rehman). Should they appear on Leadership or Supporters?"],
            ["67", "Education programmes", "Paramedical School, Nursing Teaching College, Doctor Training Program — approved copy and photos?"],
            ["68", "Patient Information Guide", "OPD / OT day charts were only available as images. Please send updated timings as text or as files we may publish."],
            ["69", "Privacy / terms / legal footer links", "Not on the site yet. Please send Privacy Policy and Terms (or say they are not required for launch)."],
            ["70", "Committee titles", "Some names still look like notes, e.g. “HR / Fundraising Marketing /& Event Management” and “Central Purchase. Machinery/ Disposable Wastage & Leakage…”. Please send the official committee names."],
            ["71", "Compliance certificates", "We have logos. If you want visitors to open a full certificate, please send a scan of each: FBR / NPO exemption, Punjab Healthcare Commission, PCP membership, ISO 9001, Shariah certificate, Social Welfare registration, Certificate of Incorporation."],
        ],
        [1.4, 5.2, 10.8],
    )

    heading(doc, "3.4 Success stories — editorial", 2)
    para(doc, "The current stories read as samples, not consented patient accounts. Please either:", space_after=4)
    bullet(doc, "supply approved stories (patient or family first name or initials, department, short quote, photo with written consent), or")
    bullet(doc, "instruct us to unpublish them until then.")

    # Section 4
    heading(doc, "4. What we already have (no action needed unless you want better)", 1)
    para(
        doc,
        "These are in place so you do not re-send them unless you have a higher-quality original:",
        space_after=6,
    )
    for item in [
        "Founder portraits: Haji Inam Elahi Asar, Mian Abdul Waheed",
        "Chairman and President portraits and messages",
        "Senior management portraits (11 people)",
        "Mrs. Tasneem Firdous Waheed portrait (biography page and dates still outstanding)",
        "Most department and service hero photos (except the stand-ins in §2)",
        "Medical Tower photo, Tamgha-e-Imtiaz image, Our Purpose / Our Impact / Supporters heroes",
        "Partner logos (Fatima Memorial, Qarshi, Hajvery, Chughtai, Al-Noor, Hormone Lab)",
        "Compliance logos (except Certificate of Incorporation)",
        "Meezan Bank logo",
        "28 of 30 donor-wall logos",
        "Donate / Patient Welfare pages using illustrations on purpose (not photos)",
    ]:
        bullet(doc, item)

    # Section 5
    heading(doc, "5. Suggested photo brief (for the photographer)", 1)
    para(doc, "If you are arranging a shoot, this is the remaining list in one place:", space_after=6)
    brief = [
        ("Portraits — ", "all names in §1.1 (passport-style, even lighting, no heavy filters)."),
        ("Clinical rooms — ", "Orthopedics, Urology, Plastic Surgery, ENT, Gastroenterology, Pulmonology, Endocrinology, Dermatology, Dietetics, Anesthesia."),
        ("Ambulance — ", "vehicle and/or bay, identifiable as Hijaz."),
        ("Cafeteria — ", "serving area and seating."),
        ("IPD rooms — ", "one photo each of Private Gynae, VIP, Semi-VIP, VVIP."),
        ("Equipment — ", "labelled close-ups for the homepage machinery strip."),
        ("Events — ", "one honest photo per news/event story still using a stand-in."),
        ("Donor wall (optional) — ", "a photograph of the physical wall at the hospital."),
    ]
    for lead, rest in brief:
        bullet(doc, rest, bold_lead=lead)
    para(
        doc,
        "Please avoid photographing identifiable patients unless a signed consent form is provided.",
        space_before=8,
        space_after=10,
        italic=True,
    )

    # Section 6
    heading(doc, "6. How to send files", 1)
    for item in [
        "Name files with the person’s or room’s name, e.g. Dr-Ayesha-Azhar.jpg, cafeteria.jpg, jazzcash-qr.png.",
        "Portraits: one person per file.",
        "Logos: on a white or transparent background if possible.",
        "For stories: photo + short approved text + consent note.",
    ]:
        bullet(doc, item)

    para(
        doc,
        "Thank you. Once this pack is received, the remaining placeholders on the public site can be replaced.",
        space_before=12,
        space_after=16,
        bold=True,
    )

    # Closing block
    close = doc.add_table(rows=3, cols=1)
    set_cell_text(close.rows[0].cells[0], "NORTHWEST INNOVATIONS  /  TwoCX LLC", bold=True, size=11, color=WHITE, fill=HEADER_FILL, align="center")
    set_cell_text(
        close.rows[1].cells[0],
        "Prepared for Hijaz Hospital Trust  ·  Website implementation",
        size=10,
        color=NAVY,
        fill=ALT_FILL,
        align="center",
    )
    set_cell_text(
        close.rows[2].cells[0],
        "Please reply with files and confirmations against item numbers 1–71.",
        size=10,
        color=MUTED,
        fill=ALT_FILL,
        align="center",
    )
    set_table_widths(close, [17.4])

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
