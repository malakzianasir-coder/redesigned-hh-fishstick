"""
Scan content/*.json for each page hero support line (excerpt, or quote/body
when that is the sole hero support), count words/characters, and write a DOCX
table report.

Run from repo root:
  python scripts/report-hero-excerpt-lengths.py
"""

from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CONTENT = ROOT / "content"
OUT = ROOT / "Hijaz-Hospital-Hero-Excerpt-Lengths.docx"

# Not loaded by the site, or not page heroes.
SKIP_FILES = {
    "departments_part1.json",
    "doctors.json",
    "lab-tests.json",
    "navigation.json",
    "site-settings.json",
    "forms.json",
    "news.json",
    "events.json",
    "success-stories.json",
}

# Slug overrides when JSON slug ≠ public path segment.
PATH_BY_SLUG = {
    "chairman": "/leadership/messages/chairman",
    "president": "/leadership/messages/president",
    "leadership-messages": "/leadership/messages",
    "patient-welfare": "/patient-welfare",
}

NAVY = RGBColor(0x1B, 0x24, 0x41)
RED = RGBColor(0xC4, 0x00, 0x12)
MUTED = RGBColor(0x5B, 0x61, 0x73)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RULE = "D8DCE6"
HEADER_FILL = "1B2441"
ALT_FILL = "F4F6FA"

WORD_RE = re.compile(r"[A-Za-z0-9]+(?:['’-][A-Za-z0-9]+)*|[^\W\d_]+", re.UNICODE)


def set_run_font(run, *, name="Calibri", size=11, bold=False, italic=False, color=NAVY):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color=RULE, sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, size=9, color=NAVY, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = 1


def count_words(text: str) -> int:
    return len(WORD_RE.findall(text or ""))


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def guess_path(source: str, slug: str, parent_slug: str | None = None, kind: str | None = None) -> str:
    """Best-effort public path for identification (not a router)."""
    if not slug:
        return "—"
    if slug in PATH_BY_SLUG and source not in {
        "services.json",
        "departments.json",
        "donations.json",
        "patient-welfare.json",
        "how-to-donate.json",
        "what-you-can-support.json",
    }:
        return PATH_BY_SLUG[slug]
    if source == "services.json":
        return f"/services/{slug}"
    if source == "departments.json":
        return f"/departments/{slug}"
    if source == "donations.json":
        if kind == "cause":
            return f"/donate/what-you-can-support/{slug}"
        return f"/donate/{slug}"
    if source == "patient-welfare.json":
        return f"/patient-welfare/{slug}"
    if source == "patient-welfare-hub.json":
        return "/patient-welfare"
    if source == "how-to-donate.json":
        if slug != "how-to-donate":
            return f"/donate/how-to-donate/{slug}"
        return "/donate/how-to-donate"
    if source == "what-you-can-support.json":
        if slug != "what-you-can-support":
            return f"/donate/what-you-can-support/{slug}"
        return "/donate/what-you-can-support"
    if source == "donate.json":
        return "/donate"
    if source == "landing-pages.json":
        return f"/l/{slug}"
    if source == "about-us.json":
        return "/about-us"
    if slug == "home":
        return "/"
    return f"/{slug}"


def walk_json_files() -> list[Path]:
    files: list[Path] = []
    for path in sorted(CONTENT.rglob("*.json")):
        if path.name in SKIP_FILES:
            continue
        files.append(path)
    return files


def collect_from_hero(
    *,
    hero: dict,
    slug: str,
    title: str,
    source: str,
    parent_slug: str | None,
    kind: str | None,
    records: list[dict],
) -> None:
    if not isinstance(hero, dict):
        return

    # Prefer excerpt; fall back to quote (Donate hub) or body (landing pages).
    if str(hero.get("excerpt") or "").strip():
        field = "excerpt"
        text = str(hero["excerpt"])
    elif str(hero.get("quote") or "").strip():
        field = "quote"
        text = str(hero["quote"])
    elif str(hero.get("body") or "").strip():
        field = "body"
        text = str(hero["body"])
    else:
        return

    text = normalize_text(text)
    records.append(
        {
            "slug": slug or "—",
            "title": title or hero.get("title") or hero.get("heading") or "—",
            "path": guess_path(source, slug, parent_slug, kind),
            "source": source,
            "field": field,
            "text": text,
            "words": count_words(text),
            "chars": len(text),
            "chars_no_spaces": len(re.sub(r"\s+", "", text)),
        }
    )


def visit(node, *, source: str, parent_slug: str | None, records: list[dict]) -> None:
    if isinstance(node, list):
        for item in node:
            visit(item, source=source, parent_slug=parent_slug, records=records)
        return
    if not isinstance(node, dict):
        return

    # Cause pages are authored in what-you-can-support.json; donations.json
    # duplicates those heroes for kind=cause and is not the live page source.
    if source == "donations.json" and node.get("kind") == "cause" and "hero" in node:
        for key, val in node.items():
            if key == "hero":
                continue
            if isinstance(val, (dict, list)):
                visit(val, source=source, parent_slug=parent_slug, records=records)
        return

    slug = str(node.get("slug") or "").strip() or (parent_slug or "")
    title = str(node.get("title") or node.get("heading") or "").strip()
    kind = str(node.get("kind") or "").strip() or None
    hero = node.get("hero")
    if isinstance(hero, dict):
        collect_from_hero(
            hero=hero,
            slug=slug or "",
            title=title or str(hero.get("title") or hero.get("heading") or ""),
            source=source,
            parent_slug=parent_slug,
            kind=kind,
            records=records,
        )

    next_parent = slug or parent_slug
    for key, val in node.items():
        if key == "hero":
            continue
        if isinstance(val, (dict, list)):
            visit(val, source=source, parent_slug=next_parent, records=records)


def collect_records() -> list[dict]:
    records: list[dict] = []

    for path in walk_json_files():
        rel = path.relative_to(ROOT).as_posix()
        source = path.name
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError as exc:
            print(f"SKIP (parse error): {rel}: {exc}")
            continue

        if source == "landing-pages.json" and isinstance(data, dict) and "pages" in data:
            visit(data["pages"], source=source, parent_slug=None, records=records)
        else:
            if isinstance(data, dict) and "hero" in data and not data.get("slug"):
                data = {**data, "slug": path.stem}
            visit(data, source=source, parent_slug=None, records=records)

    unique: list[dict] = []
    seen: set[tuple] = set()
    for rec in records:
        key = (rec["path"], rec["field"], rec["text"])
        if key in seen:
            continue
        seen.add(key)
        unique.append(rec)

    unique.sort(key=lambda r: (r["path"], r["slug"]))
    return unique


def build_docx(records: list[dict]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    # Landscape helps the text column
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    brand = doc.add_paragraph()
    r = brand.add_run("Hijaz Hospital Trust")
    set_run_font(r, size=12, bold=True, color=RED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Hero Excerpt Length Report")
    set_run_font(r, size=20, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run(
        "Word and character counts for each page hero support line "
        "(hero.excerpt, or hero.quote / hero.body when that is the sole support)."
    )
    set_run_font(r, size=11, italic=True, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    words = [r["words"] for r in records]
    avg = (sum(words) / len(words)) if words else 0
    r = meta.add_run(
        f"Generated {date.today().isoformat()} · {len(records)} pages · "
        f"avg {avg:.1f} words · min {min(words) if words else 0} · max {max(words) if words else 0}"
    )
    set_run_font(r, size=10, color=MUTED)

    headers = (
        "#",
        "Slug",
        "Path",
        "Title",
        "Field",
        "Words",
        "Chars",
        "Chars (no spaces)",
        "Hero text",
        "Source file",
    )
    table = doc.add_table(rows=1 + len(records), cols=len(headers))
    table.autofit = True

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=WHITE, fill=HEADER_FILL)

    for idx, rec in enumerate(records):
        fill = ALT_FILL if idx % 2 else None
        values = (
            str(idx + 1),
            rec["slug"],
            rec["path"],
            rec["title"],
            rec["field"],
            str(rec["words"]),
            str(rec["chars"]),
            str(rec["chars_no_spaces"]),
            rec["text"],
            rec["source"],
        )
        for c_i, val in enumerate(values):
            set_cell_text(table.rows[idx + 1].cells[c_i], val, size=8, fill=fill)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    r = note.add_run(
        "Notes: Card/link excerpts outside hero blocks are excluded. "
        "Donate hub uses a Qur’an quote (no excerpt). Landing pages use hero.body. "
        f"Skipped files: {', '.join(sorted(SKIP_FILES))}."
    )
    set_run_font(r, size=9, italic=True, color=MUTED)

    doc.save(OUT)
    return OUT


def main() -> None:
    records = collect_records()
    out = build_docx(records)

    print(f"Found {len(records)} hero support lines\n")
    print(f"{'Slug':<36} {'Field':<8} {'Words':>5} {'Chars':>5}  Path")
    print("-" * 100)
    for rec in records:
        print(
            f"{rec['slug']:<36} {rec['field']:<8} {rec['words']:>5} {rec['chars']:>5}  {rec['path']}"
        )
    print(f"\nWrote {out}")


if __name__ == "__main__":
    main()
